import os
import logging
import pdfplumber
import pypdf
import docx

logger = logging.getLogger(__name__)

class DocumentExtractor:
    """
    Service responsible for parsing and extracting raw text from PDF, DOCX, and TXT files.
    """

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        file_type = file_type.lower().replace('.', '')

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at path: {file_path}")

        if file_type == 'pdf':
            return DocumentExtractor._extract_pdf(file_path)
        elif file_type == 'docx':
            return DocumentExtractor._extract_docx(file_path)
        elif file_type in ['txt', 'text']:
            return DocumentExtractor._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_type}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        text_pages = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_pages.append(page_text.strip())
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed for {file_path}: {e}. Falling back to pypdf.")
            text_pages = []

        # If pdfplumber & pypdf return minimal text (e.g. scanned image PDF), try Gemini Multimodal OCR fallback
        extracted_content = "\n\n".join(text_pages).strip()
        if len(extracted_content) < 50:
            try:
                from ai.services import GeminiProvider
                provider = GeminiProvider()
                if provider.client:
                    logger.info(f"Attempting Gemini Multimodal OCR fallback for scanned PDF: {file_path}")
                    uploaded_file = provider.client.files.upload(file=file_path)
                    response = provider.client.models.generate_content(
                        model=provider.model,
                        contents=[uploaded_file, "Extract all raw text from this document accurately page by page without omission."]
                    )
                    ocr_text = response.text.strip()
                    if ocr_text:
                        return ocr_text
            except Exception as ocr_err:
                logger.warning(f"Gemini Multimodal OCR fallback failed for {file_path}: {ocr_err}")

        return extracted_content

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
                        
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            raise Exception(f"Failed to extract text from DOCX file: {e}")

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                continue
        
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read().strip()

